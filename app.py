import streamlit as st
from docx import Document
import io

st.title("Notas Inacistencias")
st.subheader("""Agradecimientos
             - Sergio Cortés Prado
    - Vicente Barañao Abarzúa""")
st.write("Complete los datos pedidos")

# Create input fields
input1 = st.text_input("Número de nota (EJ: 055)")
input2 = st.text_input("Nombre del Voluntario (EJ: Pedro Pérez González)")
input3 = st.text_input("Fecha a realizarse los cursos (EJ: los días viernes 5 y sábado 6 de mayo)")
input4 = st.text_input("Razón por la que el voluntario no puede asistir al curso (EJ: tiene prueba de la universidad en ese horario)")
input5 = st.text_input("Fecha (EJ: 5 de mayo)")
input6 = st.text_input("Nombre del curso (EJ: Fuego y Tácticas)")

# Use fixed document
template_path = "nota.docx"

if all([input1, input2, input3, input4, input5, input6]):
    if st.button("Generar Nota"):
        try:
            # Read the template document
            doc = Document(template_path)
            
            # Replace placeholders with user inputs
            replacements = {
                "___1___": input1,
                "___2___": input2,
                "___3___": input3,
                "___4___": input4,
                "___5___": input5,
                "___6___": input6
            }
            
            # Process each paragraph
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
            
            # Process each table
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, value)
            
            # Save the document to a bytes buffer
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            # Create download button
            st.download_button(
                label="Descargar Nota",
                data=docx_bytes,
                file_name=f"{input1}-2025 - Inasistencia Curso EBS.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except FileNotFoundError:
            st.error("Error: The template document 'nota.docx' was not found in the current directory.")
else:
    st.warning("Por favor rellenar todos los campos!") 